"""
Flask Inventory Slip Generator - Web application for generating inventory slips
from CSV and JSON data with support for Bamboo and Cultivera formats.
"""

from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, session, send_file
import os
import sys
import json
import datetime
import urllib.request
import pandas as pd
from io import BytesIO
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches
from docxcompose.composer import Composer
import configparser
import tempfile
import uuid
import re
import werkzeug.utils
from werkzeug.utils import secure_filename
import logging
import threading
import webbrowser
from src.utils.document_handler import DocumentHandler
from src.ui.app import InventorySlipGenerator

# Configure logging
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Constants
CONFIG_FILE = os.path.expanduser("~/inventory_generator_config.ini")
DEFAULT_SAVE_DIR = os.path.expanduser("~/Downloads")
APP_VERSION = "2.0.0"
UPLOAD_FOLDER = os.path.join(tempfile.gettempdir(), "inventory_generator", "uploads")
ALLOWED_EXTENSIONS = {'csv', 'json', 'docx'}
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16 MB max upload size

# Ensure directories exist
os.makedirs(DEFAULT_SAVE_DIR, exist_ok=True)
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Initialize Flask application
app = Flask(__name__)
app.secret_key = os.urandom(24)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH
app.config['SESSION_TYPE'] = 'filesystem'

# Helper function to get resource path (for templates)
def resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller"""
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

# Load configurations or create default
def load_config():
    config = configparser.ConfigParser()
    
    # Default configurations
    config['PATHS'] = {
        'template_path': os.path.join(os.path.dirname(__file__), "templates/documents/InventorySlips.docx"),
        'output_dir': DEFAULT_SAVE_DIR,
        'recent_files': '',
        'recent_urls': ''
    }
    
    config['SETTINGS'] = {
        'items_per_page': '4',
        'auto_open': 'true',
        'theme': 'dark',
        'font_size': '12'
    }
    
    # Load existing config if it exists
    if os.path.exists(CONFIG_FILE):
        config.read(CONFIG_FILE)
    else:
        # Create config file with defaults
        with open(CONFIG_FILE, 'w') as f:
            config.write(f)
    
    return config

def save_config(config):
    with open(CONFIG_FILE, 'w') as f:
        config.write(f)

# Helper to adjust font sizes after rendering
def adjust_table_font_sizes(doc_path):
    """
    Post-process a DOCX file to dynamically adjust font size inside table cells based on thresholds.
    """
    thresholds = [
        (30, 12),   # <=30 chars → 12pt
        (45, 10),   # <=45 chars → 10pt
        (60, 8),    # <=60 chars → 8pt
        (float('inf'), 7)  # >60 chars → 7pt
    ]

    def get_font_size(text_len):
        for limit, size in thresholds:
            if text_len <= limit:
                return size
        return 7  # Fallback

    doc = Document(doc_path)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if not text:
                        continue

                    # If line is Product Name (first line), force 10pt
                    if paragraph == cell.paragraphs[0]:
                        font_size = 10
                    else:
                        font_size = get_font_size(len(text))

                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)

    doc.save(doc_path)

# Open files after saving
def open_file(path):
    try:
        if sys.platform == "darwin":
            os.system(f'open "{path}"')
        elif sys.platform == "win32":
            os.startfile(path)
        else:
            os.system(f'xdg-open "{path}"')
    except Exception as e:
        print(f"Error opening file: {e}")

# Split records into chunks
def chunk_records(records, chunk_size=4):
    for i in range(0, len(records), chunk_size):
        yield records[i:i + chunk_size]

# Check if file extension is allowed
def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Process and save inventory slips
def run_full_process_inventory_slips(selected_df, config, status_callback=None, progress_callback=None):
    if selected_df.empty:
        if status_callback:
            status_callback("Error: No data selected.")
        return False, "No data selected."

    try:
        # Get settings from config
        items_per_page = int(config['SETTINGS'].get('items_per_page', '4'))
        template_path = config['PATHS'].get('template_path')
        if not template_path or not os.path.exists(template_path):
            template_path = os.path.join(os.path.dirname(__file__), "templates/documents/InventorySlips.docx")
            if not os.path.exists(template_path):
                raise ValueError(f"Template file not found at: {template_path}")
        
        if status_callback:
            status_callback("Processing data...")

        records = selected_df.to_dict(orient="records")
        pages = []

        # Process records in chunks of 4 (or configured size)
        total_chunks = (len(records) + items_per_page - 1) // items_per_page
        current_chunk = 0

        for chunk in chunk_records(records, items_per_page):
            current_chunk += 1
            if progress_callback:
                progress = (current_chunk / total_chunks) * 50
                progress_callback(int(progress))

            if status_callback:
                status_callback(f"Generating page {current_chunk} of {total_chunks}...")

            try:
                tpl = DocxTemplate(template_path)
                context = {}

                # Fill context with records - modified vendor handling
                for idx, record in enumerate(chunk, 1):
                    # Get vendor info, using full vendor name if available
                    vendor_name = record.get("Vendor", "")
                    # If vendor is in format "license - name", extract just the name
                    if " - " in vendor_name:
                        vendor_name = vendor_name.split(" - ")[1]
                    
                    context[f"Label{idx}"] = {
                        "ProductName": record.get("Product Name*", ""),
                        "Barcode": record.get("Barcode*", ""),
                        "AcceptedDate": record.get("Accepted Date", ""),
                        "QuantityReceived": record.get("Quantity Received*", ""),
                        "Vendor": vendor_name or "Unknown Vendor",  # Only use Unknown if empty
                        "ProductType": record.get("Product Type*", "")
                    }

                # Fill remaining slots with empty values
                for i in range(len(chunk) + 1, items_per_page + 1):
                    context[f"Label{i}"] = {
                        "ProductName": "",
                        "Barcode": "",
                        "AcceptedDate": "",
                        "QuantityReceived": "",
                        "Vendor": "",
                        "ProductType": ""
                    }

                # Render template with context
                tpl.render(context)
                
                # Save to BytesIO
                output = BytesIO()
                tpl.save(output)
                pages.append(Document(output))

            except Exception as e:
                raise ValueError(f"Error generating page {current_chunk}: {e}")

        if not pages:
            return False, "No documents generated."

        # Combine pages
        if status_callback:
            status_callback("Combining pages...")

        master = pages[0]
        composer = Composer(master)
        for i, doc in enumerate(pages[1:]):
            if progress_callback:
                progress = 50 + ((i + 1) / len(pages[1:])) * 40
                progress_callback(int(progress))
            composer.append(doc)

        # Save final document
        now = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        outname = f"inventory_slips_{now}.docx"
        outpath = os.path.join(config['PATHS']['output_dir'], outname)

        if status_callback:
            status_callback("Saving document...")

        master.save(outpath)

        # Adjust font sizes
        if status_callback:
            status_callback("Adjusting formatting...")
        adjust_table_font_sizes(outpath)

        if progress_callback:
            progress_callback(100)

        return True, outpath

    except Exception as e:
        if status_callback:
            status_callback(f"Error: {str(e)}")
        return False, str(e)

# Parse Bamboo transfer schema JSON
def parse_bamboo_data(json_data):
    if not json_data:
        return pd.DataFrame()
    
    try:
        # Get vendor information
        from_license_number = json_data.get("from_license_number", "")
        from_license_name = json_data.get("from_license_name", "")
        vendor_meta = f"{from_license_number} - {from_license_name}"
        
        # Get transfer date
        raw_date = json_data.get("est_arrival_at", "") or json_data.get("transferred_at", "")
        accepted_date = raw_date.split("T")[0] if "T" in raw_date else raw_date
        
        # Process inventory items
        items = json_data.get("inventory_transfer_items", [])
        records = []
        
        for item in items:
            # Extract THC and CBD content from lab_result_data if available
            thc_content = ""
            cbd_content = ""
            
            lab_data = item.get("lab_result_data", {})
            if lab_data and "potency" in lab_data:
                for potency_item in lab_data["potency"]:
                    if potency_item.get("type") == "total-thc":
                        thc_content = f"{potency_item.get('value', '')}%"
                    elif potency_item.get("type") == "total-cbd":
                        cbd_content = f"{potency_item.get('value', '')}%"
            
            records.append({
                "Product Name*": item.get("product_name", ""),
                "Product Type*": item.get("inventory_type", ""),
                "Quantity Received*": item.get("qty", ""),
                "Barcode*": item.get("inventory_id", "") or item.get("external_id", ""),
                "Accepted Date": accepted_date,
                "Vendor": vendor_meta,
                "Strain Name": item.get("strain_name", ""),
                "THC Content": thc_content,
                "CBD Content": cbd_content,
                "Source System": "Bamboo"
            })
        
        return pd.DataFrame(records)
    
    except Exception as e:
        raise ValueError(f"Failed to parse Bamboo transfer data: {e}")

# Parse Cultivera JSON
def parse_cultivera_data(json_data):
    if not json_data:
        return pd.DataFrame()
    
    try:
        # Check if Cultivera format
        if not json_data.get("data") or not isinstance(json_data.get("data"), dict):
            raise ValueError("Not a valid Cultivera format")
        
        data = json_data.get("data", {})
        manifest = data.get("manifest", {})
        
        # Get vendor information
        from_license = manifest.get("from_license", {})
        vendor_name = from_license.get("name", "")
        vendor_license = from_license.get("license_number", "")
        vendor_meta = f"{vendor_license} - {vendor_name}" if vendor_license and vendor_name else "Unknown Vendor"
        
        # Get transfer date
        created_at = manifest.get("created_at", "")
        accepted_date = created_at.split("T")[0] if "T" in created_at else created_at
        
        # Process inventory items
        items = manifest.get("items", [])
        records = []
        
        for item in items:
            # Extract product info
            product = item.get("product", {})
            
            # Extract THC and CBD content
            thc_content = ""
            cbd_content = ""
            
            test_results = item.get("test_results", [])
            if test_results:
                for result in test_results:
                    if "thc" in result.get("type", "").lower():
                        thc_content = f"{result.get('percentage', '')}%"
                    elif "cbd" in result.get("type", "").lower():
                        cbd_content = f"{result.get('percentage', '')}%"
            
            records.append({
                "Product Name*": product.get("name", ""),
                "Product Type*": product.get("category", ""),
                "Quantity Received*": item.get("quantity", ""),
                "Barcode*": item.get("barcode", "") or item.get("id", ""),
                "Accepted Date": accepted_date,
                "Vendor": vendor_meta,
                "Strain Name": product.get("strain_name", ""),
                "THC Content": thc_content,
                "CBD Content": cbd_content,
                "Source System": "Cultivera"
            })
        
        return pd.DataFrame(records)
    
    except Exception as e:
        raise ValueError(f"Failed to parse Cultivera data: {e}")

# Detect and parse JSON from multiple systems
def parse_inventory_json(json_data):
    """
    Detects the JSON format and parses it accordingly
    """
    if not json_data:
        return None, "No data provided"
    
    try:
        # If data is a string, parse it to JSON
        if isinstance(json_data, str):
            json_data = json.loads(json_data)
        
        # Try parsing as Bamboo
        if "inventory_transfer_items" in json_data:
            return parse_bamboo_data(json_data), "Bamboo"
        
        # Try parsing as Cultivera
        elif "data" in json_data and isinstance(json_data["data"], dict) and "manifest" in json_data["data"]:
            return parse_cultivera_data(json_data), "Cultivera"
        
        # Unknown format
        else:
            return None, "Unknown JSON format. Please use Bamboo or Cultivera format."
    
    except json.JSONDecodeError:
        return None, "Invalid JSON data. Please check the format."
    except Exception as e:
        return None, f"Error parsing data: {str(e)}"

# Process CSV data
def process_csv_data(df):
    try:
        logger.info(f"Original columns: {df.columns.tolist()}")
        
        # First, ensure column names are unique by adding a suffix if needed
        df.columns = [f"{col}_{i}" if df.columns.tolist().count(col) > 1 else col 
                     for i, col in enumerate(df.columns)]
        logger.info(f"Columns after ensuring uniqueness: {df.columns.tolist()}")
        
        # Map column names to expected format
        col_map = {
            "Product Name*": "Product Name*",
            "Product Name": "Product Name*",
            "Quantity Received": "Quantity Received*",
            "Quantity*": "Quantity Received*",
            "Quantity": "Quantity Received*",
            "Lot Number*": "Barcode*",
            "Barcode": "Barcode*",
            "Lot Number": "Barcode*",
            "Accepted Date": "Accepted Date",
            "Vendor": "Vendor",
            "Strain Name": "Strain Name",
            "Product Type*": "Product Type*",
            "Product Type": "Product Type*",
            "Inventory Type": "Product Type*"
        }
        
        # Now rename columns according to our mapping
        new_columns = {}
        target_counts = {}  # Keep track of how many times we've used each target name
        
        for col in df.columns:
            base_col = col.split('_')[0]  # Remove any suffix
            if base_col in col_map:
                target_name = col_map[base_col]
                # If we've seen this target name before, add a suffix
                if target_name in target_counts:
                    target_counts[target_name] += 1
                    new_columns[col] = f"{target_name}_{target_counts[target_name]}"
                else:
                    target_counts[target_name] = 0
                    new_columns[col] = target_name
            else:
                new_columns[col] = col
        
        logger.info(f"Column mapping: {new_columns}")
        df = df.rename(columns=new_columns)
        logger.info(f"Columns after renaming: {df.columns.tolist()}")
        
        # Ensure required columns exist
        required_cols = ["Product Name*", "Barcode*"]
        missing_cols = [col for col in required_cols if not any(col in c for c in df.columns)]
        
        if missing_cols:
            return None, f"CSV is missing required columns: {', '.join(missing_cols)}"
        
        # Set default values for missing columns
        if not any("Vendor" in c for c in df.columns):
            df["Vendor"] = "Unknown Vendor"
        else:
            vendor_col = next(c for c in df.columns if "Vendor" in c)
            df[vendor_col] = df[vendor_col].fillna("Unknown Vendor")
        
        if not any("Accepted Date" in c for c in df.columns):
            today = datetime.datetime.today().strftime("%Y-%m-%d")
            df["Accepted Date"] = today
        
        if not any("Product Type*" in c for c in df.columns):
            df["Product Type*"] = "Unknown"
        
        if not any("Strain Name" in c for c in df.columns):
            df["Strain Name"] = ""
        
        # Sort if possible
        try:
            sort_cols = []
            if any("Product Type*" in c for c in df.columns):
                sort_cols.append(next(c for c in df.columns if "Product Type*" in c))
            if any("Product Name*" in c for c in df.columns):
                sort_cols.append(next(c for c in df.columns if "Product Name*" in c))
            
            if sort_cols:
                df = df.sort_values(sort_cols, ascending=[True, True])
        except:
            pass  # If sorting fails, continue without sorting
        
        # Final check for duplicate columns
        if len(df.columns) != len(set(df.columns)):
            duplicates = [col for col in df.columns if df.columns.tolist().count(col) > 1]
            logger.error(f"Duplicate columns found: {duplicates}")
            return None, f"Duplicate columns found: {', '.join(duplicates)}"
        
        return df, "Success"
    
    except Exception as e:
        logger.error(f"Error in process_csv_data: {str(e)}", exc_info=True)
        return None, f"Failed to process CSV data: {e}"

# Flask Routes
@app.route('/')
def index():
    # Load configuration
    config = load_config()
    
    # Load any previously saved data from session
    df_json = session.get('df_json', None)
    format_type = session.get('format_type', None)
    
    return render_template(
        'index.html',
        version=APP_VERSION,
        theme=config['SETTINGS'].get('theme', 'dark'),
        df_json=df_json,
        format_type=format_type,
        config=config  # Pass the application config instead of Flask config
    )

@app.route('/upload-csv', methods=['POST'])
def upload_csv():
    if 'file' not in request.files:
        logger.error('No file part in request')
        flash('No file part')
        return redirect(request.url)
    
    file = request.files['file']
    logger.info(f'Received file: {file.filename}')
    
    if file.filename == '':
        logger.error('No selected file')
        flash('No selected file')
        return redirect(request.url)
    
    if file and allowed_file(file.filename):
        try:
            # Read CSV
            logger.info('Reading CSV file')
            df = pd.read_csv(file)
            logger.info(f'CSV columns: {df.columns.tolist()}')
            
            # Process CSV
            logger.info('Processing CSV data')
            result_df, message = process_csv_data(df)
            
            if result_df is None:
                logger.error(f'Error processing CSV: {message}')
                flash(f'Error: {message}')
                return redirect(url_for('index'))
            
            # Store in session (convert DataFrame to JSON for storage)
            logger.info('Converting DataFrame to JSON')
            session['df_json'] = result_df.to_json(orient='records')
            session['format_type'] = 'CSV'
            
            flash('CSV file processed successfully')
            return redirect(url_for('data_view'))
        
        except Exception as e:
            logger.error(f'Error processing file: {str(e)}', exc_info=True)
            flash(f'Error processing file: {str(e)}')
            return redirect(url_for('index'))
    
    logger.error(f'Invalid file type: {file.filename}')
    flash('Invalid file type. Please upload a CSV file')
    return redirect(url_for('index'))

@app.route('/upload-json', methods=['POST'])
def upload_json():
    if 'file' not in request.files:
        flash('No file part')
        return redirect(request.url)
    
    file = request.files['file']
    
    if file.filename == '':
        flash('No selected file')
        return redirect(request.url)
    
    if file and allowed_file(file.filename):
        try:
            # Read JSON
            json_data = json.load(file)
            
            # Process JSON
            result_df, format_type = parse_inventory_json(json_data)
            
            if result_df is None:
                flash(f'Error: {format_type}')
                return redirect(url_for('index'))
            
            # Store in session
            session['df_json'] = result_df.to_json(orient='records')
            session['format_type'] = format_type
            session['raw_json'] = json.dumps(json_data)
            
            flash(f'{format_type} data processed successfully')
            return redirect(url_for('data_view'))
        
        except Exception as e:
            flash(f'Error processing JSON file: {str(e)}')
            return redirect(url_for('index'))
    
    flash('Invalid file type. Please upload a JSON file')
    return redirect(url_for('index'))

@app.route('/paste-json', methods=['POST'])
def paste_json():
    try:
        json_data = request.form.get('json_text', '')
        api_format = request.form.get('api_format', 'auto')
        
        # Process JSON based on format
        if api_format == 'bamboo':
            data = json.loads(json_data)
            result_df = parse_bamboo_data(data)
            format_type = 'Bamboo'
        elif api_format == 'cultivera':
            data = json.loads(json_data)
            result_df = parse_cultivera_data(data)
            format_type = 'Cultivera'
        else:
            # Auto-detect format
            result_df, format_type = parse_inventory_json(json_data)
        
        if result_df is None or result_df.empty:
            flash(f'Could not process {api_format} data')
            return redirect(url_for('index'))
        
        # Store in session
        session['df_json'] = result_df.to_json(orient='records')
        session['format_type'] = format_type
        session['raw_json'] = json_data
        
        flash(f'{format_type} data imported successfully')
        return redirect(url_for('data_view'))
    
    except Exception as e:
        flash(f'Failed to import data: {str(e)}')
        return redirect(url_for('index'))

@app.route('/load-url', methods=['POST'])
def load_url():
    url = request.form.get('url', '')
    
    if not url.startswith('http'):
        flash('Please enter a valid URL starting with http:// or https://')
        return redirect(url_for('index'))
    
    try:
        # Fetch data from URL
        with urllib.request.urlopen(url) as resp:
            data = json.loads(resp.read().decode())
        
        # Process the data
        result_df, format_type = parse_inventory_json(data)
        
        if result_df is None:
            flash(f'Could not parse data: {format_type}')
            return redirect(url_for('index'))
        
        # Store in session
        session['df_json'] = result_df.to_json(orient='records')
        session['format_type'] = format_type
        session['raw_json'] = json.dumps(data)
        
        # Add to recent URLs
        config = load_config()
        recent_urls = config['PATHS'].get('recent_urls', '').split('|')
        recent_urls = [u for u in recent_urls if u]
        if url not in recent_urls:
            recent_urls.insert(0, url)
            recent_urls = recent_urls[:10]  # Keep only 10 most recent
            config['PATHS']['recent_urls'] = '|'.join(recent_urls)
            save_config(config)
        
        flash(f'{format_type} data loaded successfully from URL')
        return redirect(url_for('data_view'))
    
    except Exception as e:
        flash(f'Failed to load data from URL: {str(e)}')
        return redirect(url_for('index'))

@app.route('/fetch-api', methods=['POST'])
def fetch_api():
    url = request.form.get('url', '')
    api_type = request.form.get('api_type', 'auto')
    api_key = request.form.get('api_key', '')
    
    if not url:
        flash('Please enter an API URL')
        return redirect(url_for('index'))
    
    try:
        # Set up headers
        headers = {
            "User-Agent": "InventorySlipGenerator/2.0.0",
            "Accept": "application/json",
            "Content-Type": "application/json"
        }
        
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
        
        # Save API key to config
        config = load_config()
        if 'API' not in config:
            config['API'] = {}
        config['API']['bamboo_key'] = api_key
        save_config(config)
        
        # Fetch data
        req = urllib.request.Request(url, headers=headers)
        with urllib.request.urlopen(req) as resp:
            data = json.loads(resp.read().decode())
        
        # Process based on API type
        if api_type == 'bamboo':
            result_df = parse_bamboo_data(data)
            format_type = 'Bamboo'
        elif api_type == 'cultivera':
            result_df = parse_cultivera_data(data)
            format_type = 'Cultivera'
        else:
            # Auto-detect
            result_df, format_type = parse_inventory_json(data)
        
        if result_df is None or result_df.empty:
            flash(f'Could not process {api_type} data')
            return redirect(url_for('index'))
        
        # Store in session
        session['df_json'] = result_df.to_json(orient='records')
        session['format_type'] = format_type
        session['raw_json'] = json.dumps(data)
        
        # Add to recent URLs
        recent_urls = config['PATHS'].get('recent_urls', '').split('|')
        recent_urls = [u for u in recent_urls if u]
        if url not in recent_urls:
            recent_urls.insert(0, url)
            recent_urls = recent_urls[:10]  # Keep only 10 most recent
            config['PATHS']['recent_urls'] = '|'.join(recent_urls)
            save_config(config)
        
        flash(f'{format_type} data fetched successfully from API')
        return redirect(url_for('data_view'))
    
    except Exception as e:
        flash(f'Failed to fetch API data: {str(e)}')
        return redirect(url_for('index'))

@app.route('/data-view')
def data_view():
    # Load data from session
    df_json = session.get('df_json', None)
    format_type = session.get('format_type', None)
    
    if df_json is None:
        flash('No data available. Please load data first.')
        return redirect(url_for('index'))
    
    # Convert JSON to DataFrame
    df = pd.read_json(df_json, orient='records')
    
    # Format data for template
    products = []
    for idx, row in df.iterrows():
        product = {
            'id': idx,
            'name': row.get('Product Name*', ''),
            'strain': row.get('Strain Name', ''),
            'sku': row.get('Barcode*', ''),
            'quantity': row.get('Quantity Received*', ''),
            'source': format_type or 'Unknown'
        }
        products.append(product)
    
    # Load configuration
    config = load_config()
    
    return render_template(
        'data_view.html',
        products=products,
        format_type=format_type,
        theme=config['SETTINGS'].get('theme', 'dark'),
        version=APP_VERSION
    )

@app.route('/generate-slips', methods=['POST'])
def generate_slips():
    try:
        # Get selected products
        selected_indices = request.form.getlist('selected_indices[]')
        
        if not selected_indices:
            flash('No products selected.')
            return redirect(url_for('data_view'))
        
        # Convert indices to integers
        selected_indices = [int(idx) for idx in selected_indices]
        
        # Load data from session
        df_json = session.get('df_json', None)
        
        if df_json is None:
            flash('No data available. Please load data first.')
            return redirect(url_for('index'))
        
        # Convert JSON to DataFrame
        df = pd.read_json(df_json, orient='records')
        
        # Get only selected rows
        selected_df = df.iloc[selected_indices].copy()
        
        # Load configuration
        config = load_config()
        
        # Generate the file
        status_messages = []
        progress_values = []
        
        def status_callback(msg):
            status_messages.append(msg)
        
        def progress_callback(value):
            progress_values.append(value)
        
        success, result = run_full_process_inventory_slips(
            selected_df,
            config,
            status_callback,
            progress_callback
        )
        
        if success:
            # Return the file for download
            return send_file(
                result,
                as_attachment=True,
                download_name=os.path.basename(result),
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            flash(f'Failed to generate inventory slips: {result}')
            return redirect(url_for('data_view'))
    
    except Exception as e:
        flash(f'Error generating slips: {str(e)}')
        return redirect(url_for('data_view'))

@app.route('/show-result')
def show_result():
    # Get output file path from session
    output_file = session.get('output_file', None)
    
    if not output_file or not os.path.exists(output_file):
        flash('No output file available.')
        return redirect(url_for('index'))
    
    # Get filename for display
    filename = os.path.basename(output_file)
    
    # Load configuration
    config = load_config()
    
    return render_template(
        'result.html',
        filename=filename,
        theme=config['SETTINGS'].get('theme', 'dark'),
        version=APP_VERSION
    )

@app.route('/download-file')
def download_file():
    # Get output file path from session
    output_file = session.get('output_file', None)
    
    if not output_file or not os.path.exists(output_file):
        flash('No output file available.')
        return redirect(url_for('index'))
    
    # Return the file for download
    return send_file(output_file, as_attachment=True)

@app.route('/settings', methods=['GET', 'POST'])
def settings():
    config = load_config()
    
    if request.method == 'POST':
        # Update settings from form
        if 'items_per_page' in request.form:
            config['SETTINGS']['items_per_page'] = request.form['items_per_page']
        
        if 'theme' in request.form:
            config['SETTINGS']['theme'] = request.form['theme']
        
        if 'api_key' in request.form:
            if 'API' not in config:
                config['API'] = {}
            config['API']['bamboo_key'] = request.form['api_key']
        
        if 'output_dir' in request.form:
            output_dir = request.form['output_dir']
            if output_dir and os.path.exists(output_dir):
                config['PATHS']['output_dir'] = output_dir
        
        # Save updated config
        save_config(config)
        flash('Settings saved successfully')
        return redirect(url_for('index'))
    
    return render_template(
        'settings.html',
        config=config,
        theme=config['SETTINGS'].get('theme', 'dark'),
        version=APP_VERSION
    )

@app.route('/api-import')
def api_import():
    config = load_config()
    
    # Get Bamboo API key from config
    api_key = ''
    if 'API' in config and 'bamboo_key' in config['API']:
        api_key = config['API']['bamboo_key']
    
    # Get recent URLs
    recent_urls = config['PATHS'].get('recent_urls', '').split('|')
    recent_urls = [u for u in recent_urls if u]
    
    return render_template(
        'api_import.html',
        api_key=api_key,
        recent_urls=recent_urls,
        theme=config['SETTINGS'].get('theme', 'dark'),
        version=APP_VERSION
    )

@app.route('/view-json')
def view_json():
    raw_json = session.get('raw_json', None)
    format_type = session.get('format_type', None)
    
    if raw_json is None:
        flash('No JSON data available.')
        return redirect(url_for('index'))
    
    # Load configuration
    config = load_config()
    
    return render_template(
        'view_json.html',
        raw_json=raw_json,
        format_type=format_type,
        theme=config['SETTINGS'].get('theme', 'dark'),
        version=APP_VERSION
    )

@app.route('/clear-data')
def clear_data():
    # Clear session data
    session.pop('df_json', None)
    session.pop('format_type', None)
    session.pop('raw_json', None)
    session.pop('output_file', None)
    
    flash('Data cleared successfully')
    return redirect(url_for('index'))

@app.route('/about')
def about():
    config = load_config()
    return render_template(
        'about.html',
        version=APP_VERSION,
        theme=config['SETTINGS'].get('theme', 'dark')
    )

# Error handlers
@app.errorhandler(404)
def page_not_found(e):
    config = load_config()
    return render_template('404.html', theme=config['SETTINGS'].get('theme', 'dark')), 404

@app.errorhandler(500)
def server_error(e):
    config = load_config()
    return render_template('500.html', theme=config['SETTINGS'].get('theme', 'dark')), 500

def validate_docx(file_path):
    """Validate the generated DOCX file"""
    try:
        doc = Document(file_path)
        # Try to access content to verify document is readable
        _ = doc.paragraphs
        _ = doc.tables
        return True
    except Exception as e:
        logger.error(f"Document validation failed: {str(e)}")
        return False

if __name__ == '__main__':
    # Ensure template directories exist
    os.makedirs('templates', exist_ok=True)
    
    # Create templates directory structure if it doesn't exist
    template_dirs = [
        'templates',
        'templates/documents'
    ]
    
    for directory in template_dirs:
        os.makedirs(directory, exist_ok=True)
    
    # Create default template if it doesn't exist
    default_template = os.path.join(os.path.dirname(__file__), "templates/documents/InventorySlips.docx")
    if not os.path.exists(default_template):
        doc = Document()
        # Set up template formatting
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        
        # Save template
        doc.save(default_template)
        logger.info(f"Created default template at: {default_template}")
    
    app.run(debug=True, host='0.0.0.0', port=8000)